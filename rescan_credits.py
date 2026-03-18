#!/usr/bin/env python3
"""
Rescan all DOCX credit blocks, compare with existing table,
and generate a new Excel with highlighting:
  - Light green: new rows/cells
  - Light red: changed cells
"""

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import os
import re
from datetime import datetime

# ── Configuration ──
OLD_EXCEL = "tdw-creditroll-2026-03-08.xlsx"
TODAY = datetime.now().strftime("%Y-%m-%d")
NEW_EXCEL = f"tdw-creditroll-{TODAY}.xlsx"
DOCX_DIR = "in"

FILL_GREEN = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")  # light green
FILL_RED = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")    # light red

# Keywords for Firma detection
FIRMA_KEYWORDS = ['unterstützt', 'in zusammenarbeit', 'koproduktion', 'in auftrag',
                  'gefördert', 'partner', 'förderung']


def extract_id(filename):
    m = re.match(r'(A_\d+)', filename)
    return m.group(1) if m else None


def cell_has_green(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.highlight_color and 'GREEN' in str(run.font.highlight_color):
                return True
    return False


def is_firma_role(role):
    lower = role.lower()
    return any(kw in lower for kw in FIRMA_KEYWORDS)


def split_names(text):
    parts = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        for name in line.split(','):
            name = name.strip()
            if name:
                parts.append(name)
    return parts


def scan_docx(filepath):
    """Extract credits, production title, curator from a DOCX."""
    doc = Document(filepath)
    entries = []
    prod_title = ''
    curator = ''

    # Production title from Table 0
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        if len(t0.rows) > 0 and len(t0.columns) > 1:
            prod_title = t0.rows[0].cells[1].text.strip()

    # Curator from Table 1
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            if 'kurator' in row.cells[0].text.strip().lower():
                if len(row.cells) > 1:
                    curator = row.cells[1].text.strip()
                break

    # Credits from 2-column table
    credit_table = None
    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) >= 3:
            first = table.rows[0].cells[0].text.strip().lower()
            if first in ('medium / datei', 'medium', 'datei'):
                continue
            credit_table = table
            break

    if credit_table:
        for row in credit_table.rows:
            role = row.cells[0].text.strip()
            names_text = row.cells[1].text.strip()
            if not names_text:
                continue

            is_green = cell_has_green(row.cells[0]) or cell_has_green(row.cells[1])
            markiert = 'Ja' if is_green else 'Nein'
            firma = is_firma_role(role)

            for name in split_names(names_text):
                name_clean = re.sub(r'\(.*?\)', '', name).strip()
                if not name_clean:
                    continue

                typ = 'Firma' if firma else 'Mensch'
                entries.append({
                    'Name': name_clean,
                    'Rolle': role,
                    'Markiert': markiert,
                    'Typ': typ
                })

    return entries, prod_title, curator


def main():
    # 1. Read existing table
    print(f"Reading old table: {OLD_EXCEL}")
    df_old = pd.read_excel(OLD_EXCEL)
    old_cols = list(df_old.columns)
    print(f"  Old: {len(df_old)} rows")

    # Create lookup by Name for comparison
    old_lookup = {}
    for idx, row in df_old.iterrows():
        old_lookup[str(row['Name']).strip()] = row.to_dict()

    # 2. Scan all DOCX files
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx') and not f.startswith('.')])
    print(f"Scanning {len(docx_files)} DOCX files...")

    all_entries = []
    for fname in docx_files:
        pid = extract_id(fname)
        if not pid:
            continue

        entries, prod_title, curator = scan_docx(os.path.join(DOCX_DIR, fname))
        print(f"  {fname}: {len(entries)} credits (title: '{prod_title[:30]}')")

        # Get base info
        base = {}
        match = df_old[df_old['ID'] == pid]
        if len(match) > 0:
            sample = match.iloc[0]
            for col in ['Art', 'Zielgruppe', 'Chronologische_Sortierung', 'Herkunftsland']:
                base[col] = sample.get(col, '')
            base['Titel'] = sample.get('Titel', prod_title)
        else:
            base['Titel'] = prod_title

        for entry in entries:
            entry['ID'] = pid
            entry.update(base)
        all_entries.extend(entries)

        # Add title row
        if prod_title:
            all_entries.append({
                'Name': prod_title, 'Rolle': 'Titel', 'Markiert': 'Ja',
                'Typ': 'Titel', 'ID': pid, **base
            })

        # Add curator
        if curator:
            all_entries.append({
                'Name': curator, 'Rolle': 'Kurator', 'Markiert': 'Ja',
                'Typ': 'Mensch', 'ID': pid, **base
            })

    print(f"Total extracted: {len(all_entries)}")

    # 3. Build new DataFrame, group by Name
    df_new_raw = pd.DataFrame(all_entries)

    results = []
    for name, group in df_new_raw.groupby('Name', sort=False):
        all_roles = []
        for val in group['Rolle'].dropna().unique():
            v = str(val).strip()
            if v and v != 'nan' and v not in all_roles:
                all_roles.append(v)

        base = group.iloc[0].copy()
        for i in range(4):
            base[f'Rolle {i+1}'] = all_roles[i] if i < len(all_roles) else None
        if 'Ja' in group['Markiert'].values:
            base['Markiert'] = 'Ja'

        typs = set(group['Typ'].dropna().values)
        if 'Titel' in typs:
            base['Typ'] = 'Titel'
        elif 'Firma' in typs:
            base['Typ'] = 'Firma'
        else:
            base['Typ'] = 'Mensch'

        results.append(base)

    df_scanned = pd.DataFrame(results)
    if 'Rolle' in df_scanned.columns:
        df_scanned.drop(columns=['Rolle'], inplace=True)

    # 4. Merge scanned data into old table
    # Start with old table, update and add
    df_final = df_old.copy()

    # Track changes for highlighting
    changed_cells = {}  # (row_idx, col_name) -> 'changed' or 'new'
    new_row_indices = set()

    existing_names = set(df_final['Name'].astype(str).str.strip())

    for _, new_row in df_scanned.iterrows():
        name = str(new_row['Name']).strip()

        if name in existing_names:
            # Update existing row
            idx = df_final[df_final['Name'].astype(str).str.strip() == name].index[0]
            for col in ['Markiert', 'Typ', 'Rolle 1', 'Rolle 2', 'Rolle 3', 'Rolle 4']:
                if col not in df_final.columns:
                    continue
                old_val = str(df_final.loc[idx, col]).strip() if pd.notna(df_final.loc[idx, col]) else ''
                new_val = str(new_row.get(col, '')).strip() if pd.notna(new_row.get(col, '')) else ''
                if old_val != new_val and new_val:
                    df_final.loc[idx, col] = new_row.get(col)
                    changed_cells[(idx, col)] = 'changed'
        else:
            # New entry
            new_entry = {}
            for col in df_final.columns:
                new_entry[col] = new_row.get(col, '')
            new_idx = len(df_final)
            df_final = pd.concat([df_final, pd.DataFrame([new_entry])], ignore_index=True)
            new_row_indices.add(new_idx)
            existing_names.add(name)

    # Sort alphabetically
    df_final.sort_values('Name', inplace=True, key=lambda x: x.astype(str).str.lower())
    df_final.reset_index(drop=True, inplace=True)

    # Build index mapping (old idx -> new idx after sort)
    name_to_new_idx = {}
    for idx, row in df_final.iterrows():
        name_to_new_idx[str(row['Name']).strip()] = idx

    print(f"\nResult: {len(df_final)} rows")
    print(f"  Changed cells: {len(changed_cells)}")
    print(f"  New rows: {len(new_row_indices)}")
    print(f"  Typ: {df_final['Typ'].value_counts().to_dict()}")

    # 5. Save to Excel
    df_final.to_excel(NEW_EXCEL, index=False)

    # 6. Apply highlighting with openpyxl
    wb = load_workbook(NEW_EXCEL)
    ws = wb.active
    col_names = list(df_final.columns)

    # Map column names to Excel column indices (1-based)
    col_map = {name: i + 1 for i, name in enumerate(col_names)}

    # Find which names are in new rows
    new_row_names = set()
    for old_idx in new_row_indices:
        # Find the name that was at this old index
        # We need to map back through the sort
        pass

    # Simpler approach: re-identify new and changed rows by name
    old_names = set(df_old['Name'].astype(str).str.strip())

    for excel_row in range(2, len(df_final) + 2):  # Excel rows (1-based, row 1 = header)
        df_row = excel_row - 2  # DataFrame index
        name = str(df_final.iloc[df_row]['Name']).strip()

        if name not in old_names:
            # New row → highlight entire row green
            for col_idx in range(1, len(col_names) + 1):
                ws.cell(row=excel_row, column=col_idx).fill = FILL_GREEN
        else:
            # Check cell-level changes
            old_row = old_lookup.get(name, {})
            for col_name in ['Markiert', 'Typ', 'Rolle 1', 'Rolle 2', 'Rolle 3', 'Rolle 4']:
                if col_name not in col_map:
                    continue
                old_val = str(old_row.get(col_name, '')).strip() if pd.notna(old_row.get(col_name, '')) else ''
                new_val = str(df_final.iloc[df_row].get(col_name, '')).strip() if pd.notna(df_final.iloc[df_row].get(col_name, '')) else ''
                if old_val != new_val and new_val:
                    ws.cell(row=excel_row, column=col_map[col_name]).fill = FILL_RED

    wb.save(NEW_EXCEL)
    print(f"\nSaved with highlighting: {NEW_EXCEL}")


if __name__ == "__main__":
    main()
