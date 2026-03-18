#!/usr/bin/env python3
"""
Scan DOCX files in 'in/' directory, extract credit names with green highlighting,
and merge into the base Excel table.

Logic:
- Credits are in 2-column tables (role | names) inside each DOCX.
- Green highlighted cells → Markiert = Ja
- Non-highlighted cells → Markiert = Nein
- Names in column 1 are split by comma or newline.
- Names from column 0 "Mit" / "Von und mit" rows are also extracted.
- Each name gets the role from column 0, the project ID from the filename.
- The final table has no duplicate names (per project), and 'Markiert' is Ja/Nein.
"""

import pandas as pd
from docx import Document
import os
import re

INPUT_EXCEL = "tdw-creditroll-2026-03-03.xlsx"
OUTPUT_EXCEL = "tdw-creditroll-2026-03-03.xlsx"  # Overwrite in-place
DOCX_DIR = "in"


def extract_id_from_filename(filename):
    """Extract project ID like 'A_003' from filename."""
    match = re.match(r'(A_\d+)', filename)
    return match.group(1) if match else None


def cell_has_green(cell):
    """Check if any run in a cell has BRIGHT_GREEN highlighting."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.highlight_color and 'GREEN' in str(run.font.highlight_color):
                return True
    return False


def split_names(names_text):
    """Split a names string into individual names."""
    # First split by newline, then by comma
    parts = []
    for line in names_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Split by comma
        for name in line.split(','):
            name = name.strip()
            if name:
                parts.append(name)
    return parts


def is_unterstützt_role(role):
    """Check if a role is a 'support/partner' type role."""
    lower = role.lower()
    return any(kw in lower for kw in ['unterstützt', 'in zusammenarbeit', 'koproduktion',
                                        'in auftrag', 'gefördert', 'partner'])


def extract_credits_from_docx(filepath):
    """
    Extract credit entries from a DOCX file.
    Returns: (entries_list, production_title, curator_name)
    """
    doc = Document(filepath)
    entries = []
    production_title = ''
    curator_name = ''

    # Extract production title from Table 0, Row 0, Col 1
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        if len(table0.rows) > 0 and len(table0.columns) > 1:
            production_title = table0.rows[0].cells[1].text.strip()

    # Extract Kurator:in from Table 1 (Faktenblock)
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        for row in table1.rows:
            if 'kurator' in row.cells[0].text.strip().lower():
                if len(row.cells) > 1:
                    curator_name = row.cells[1].text.strip()
                break

    # Find the credit table: 2-column table that ISN'T the media table
    credit_table = None
    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) >= 3:
            first_cell = table.rows[0].cells[0].text.strip().lower()
            if first_cell in ('medium / datei', 'medium', 'datei'):
                continue
            credit_table = table
            break

    if credit_table is None:
        return entries, production_title, curator_name

    for row in credit_table.rows:
        role_cell = row.cells[0]
        names_cell = row.cells[1]

        role = role_cell.text.strip()
        names_text = names_cell.text.strip()

        if not names_text:
            continue

        # Determine if green highlighted
        is_green = cell_has_green(role_cell) or cell_has_green(names_cell)

        # Determine Markiert and Firma values
        markiert = 'Ja' if is_green else 'Nein'
        firma = 'Ja' if is_unterstützt_role(role) else 'Nein'

        names = split_names(names_text)

        for name in names:
            name_clean = re.sub(r'\(.*?\)', '', name).strip()
            if not name_clean:
                continue
            entries.append({
                'Rolle': role,
                'Name': name_clean,
                'Markiert': markiert,
                'Firma': firma
            })

    return entries, production_title, curator_name


def main():
    # 1. Read existing base table
    print(f"Reading base table: {INPUT_EXCEL}")
    df_base = pd.read_excel(INPUT_EXCEL)
    print(f"  Base table: {len(df_base)} rows, columns: {list(df_base.columns)}")

    # 2. Scan all DOCX files
    docx_files = sorted([f for f in os.listdir(DOCX_DIR) if f.endswith('.docx') and not f.startswith('.')])
    print(f"Found {len(docx_files)} DOCX files in '{DOCX_DIR}/'")

    all_new_entries = []

    for fname in docx_files:
        project_id = extract_id_from_filename(fname)
        if not project_id:
            print(f"  Skipping {fname} (no project ID)")
            continue

        filepath = os.path.join(DOCX_DIR, fname)
        entries, prod_title, curator = extract_credits_from_docx(filepath)
        print(f"  {fname}: {len(entries)} credits (ID: {project_id}, title: '{prod_title[:30]}', curator: '{curator}')")

        # Get base info from existing table or DOCX
        titel_match = df_base[df_base['ID'] == project_id]['Titel'].values
        titel = titel_match[0] if len(titel_match) > 0 else prod_title
        base_fields = {}
        for col in ['Art', 'Zielgruppe', 'Chronologische_Sortierung', 'Herkunftsland']:
            col_match = df_base[df_base['ID'] == project_id][col].values
            base_fields[col] = col_match[0] if len(col_match) > 0 else ''

        for entry in entries:
            entry['ID'] = project_id
            entry['Titel'] = titel
            entry.update(base_fields)

        all_new_entries.extend(entries)

        # Add production title row (Rolle = 'Titel', Name = production title)
        if prod_title:
            all_new_entries.append({
                'Rolle': 'Titel',
                'Name': prod_title,
                'Markiert': 'Ja',
                'Firma': 'Nein',
                'ID': project_id,
                'Titel': titel,
                **base_fields
            })

        # Add curator row (Rolle = 'Kurator', Name = curator name)
        if curator:
            all_new_entries.append({
                'Rolle': 'Kurator',
                'Name': curator,
                'Markiert': 'Ja',
                'Firma': 'Nein',
                'ID': project_id,
                'Titel': titel,
                **base_fields
            })

    print(f"\nTotal new entries extracted from DOCX: {len(all_new_entries)}")

    # 3. Create DataFrame from new entries
    df_new = pd.DataFrame(all_new_entries)

    # Map roles to Rolle 1-4 columns
    # Group by Name + ID to consolidate roles
    if len(df_new) > 0:
        grouped = df_new.groupby(['Name', 'ID']).agg({
            'Rolle': list,
            'Markiert': lambda x: 'Ja' if 'Ja' in x.values else 'Nein',
            'Firma': lambda x: 'Ja' if 'Ja' in x.values else 'Nein',
            'Titel': 'first',
            'Art': 'first',
            'Zielgruppe': 'first',
            'Chronologische_Sortierung': 'first',
            'Herkunftsland': 'first'
        }).reset_index()

        # Split roles into Rolle 1-4
        for i in range(4):
            col_name = f'Rolle {i+1}'
            grouped[col_name] = grouped['Rolle'].apply(
                lambda roles: roles[i] if i < len(roles) else None
            )
        grouped.drop(columns=['Rolle'], inplace=True)

        # Reorder to match base table
        target_cols = ['Markiert', 'Firma', 'Name', 'Rolle 1', 'Rolle 2', 'Rolle 3', 'Rolle 4',
                       'ID', 'Titel', 'Art', 'Zielgruppe', 'Chronologische_Sortierung', 'Herkunftsland']
        for col in target_cols:
            if col not in grouped.columns:
                grouped[col] = ''
        grouped = grouped[target_cols]
    else:
        grouped = pd.DataFrame(columns=df_base.columns)

    # 4. Merge: update existing rows and add new ones
    # Key for matching: Name + ID (a person in a specific project)
    df_base['_key'] = df_base['Name'].astype(str).str.strip() + '|||' + df_base['ID'].astype(str).str.strip()
    grouped['_key'] = grouped['Name'].astype(str).str.strip() + '|||' + grouped['ID'].astype(str).str.strip()

    # Update Markiert for existing entries
    existing_keys = set(df_base['_key'])
    new_keys = set(grouped['_key'])

    updated_count = 0
    added_count = 0

    for _, row in grouped.iterrows():
        key = row['_key']
        if key in existing_keys:
            # Update Markiert status
            mask = df_base['_key'] == key
            old_markiert = df_base.loc[mask, 'Markiert'].values[0]
            new_markiert = row['Markiert']
            if old_markiert != new_markiert:
                df_base.loc[mask, 'Markiert'] = new_markiert
                updated_count += 1
        else:
            # Add new entry
            new_row = row.drop('_key')
            df_base = pd.concat([df_base, pd.DataFrame([new_row])], ignore_index=True)
            added_count += 1

    # Clean up
    df_base.drop(columns=['_key'], inplace=True)

    # 5. Remove duplicates by Name (keep first occurrence, prefer Ja)
    # Sort so 'Ja' comes first, then drop duplicates
    df_base['_sort'] = df_base['Markiert'].apply(lambda x: 0 if x == 'Ja' else 1)
    df_base.sort_values('_sort', inplace=True)
    before_dedup = len(df_base)
    df_base.drop_duplicates(subset=['Name'], keep='first', inplace=True)
    df_base.drop(columns=['_sort'], inplace=True)
    dedup_removed = before_dedup - len(df_base)

    # Sort alphabetically by Name
    df_base.sort_values('Name', inplace=True, key=lambda x: x.str.lower())
    df_base.reset_index(drop=True, inplace=True)

    print(f"\nMerge results:")
    print(f"  Updated Markiert: {updated_count}")
    print(f"  Added new entries: {added_count}")
    print(f"  Duplicates removed: {dedup_removed}")
    print(f"  Final row count: {len(df_base)}")
    print(f"  Markiert counts: {df_base['Markiert'].value_counts().to_dict()}")

    # 6. Save
    df_base.to_excel(OUTPUT_EXCEL, index=False)
    print(f"\nSaved to: {OUTPUT_EXCEL}")


if __name__ == "__main__":
    main()
