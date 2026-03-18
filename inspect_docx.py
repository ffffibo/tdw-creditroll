from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import sys

doc_path = sys.argv[1]
doc = Document(doc_path)

print(f"--- paragraphs ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P: {p.text}")

print(f"--- tables ---")
for i, t in enumerate(doc.tables):
    print(f"Table {i}")
    for r, row in enumerate(t.rows):
        cells = []
        for c in row.cells:
            cell_text = ""
            for p in c.paragraphs:
                for run in p.runs:
                    color = run.font.color.rgb if run.font.color and run.font.color.rgb else "None"
                    highlight = run.font.highlight_color if run.font.highlight_color else "None"
                    cell_text += f"[{run.text}(c:{color},h:{highlight})]"
            cells.append(cell_text)
        print(f"  Row {r}: {cells}")

